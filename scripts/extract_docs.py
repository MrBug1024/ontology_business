"""提取 E:\\gx\\new_docs 下 docx/pdf 文档文本，输出到 stdout 供阅读。"""
import os
import sys

from docx import Document

BASE = r"E:\gx\new_docs"


def extract_docx(p: str) -> str:
    d = Document(p)
    parts = [par.text for par in d.paragraphs if par.text.strip()]
    for t in d.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def main() -> None:
    only = sys.argv[1] if len(sys.argv) > 1 else ""
    limit = int(sys.argv[2]) if len(sys.argv) > 2 else 4000
    for root, dirs, files in os.walk(BASE):
        for f in sorted(files):
            if f.startswith("~") or f.startswith(".~"):
                continue
            p = os.path.join(root, f)
            if not f.lower().endswith(".docx"):
                continue
            if only and only.lower() not in f.lower():
                continue
            try:
                text = extract_docx(p)
            except Exception as e:  # noqa: BLE001
                print(f"ERR {p}: {e}")
                continue
            print("=" * 90)
            print("FILE:", p)
            print("LEN:", len(text))
            print("=" * 90)
            print(text[:limit])
            print()


if __name__ == "__main__":
    main()

# -*- coding: utf-8 -*-
import os
import openpyxl
try:
    import docx
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

root = r'E:\gx\new_docs\逻辑资料V1\AI智能体相关法律、准则、底稿、案例、报告模版【20260703】'

def read_docx(p, limit=4000):
    d = docx.Document(p)
    out = []
    for para in d.paragraphs:
        t = para.text.strip()
        if t:
            out.append(t)
    # tables
    for tb in d.tables:
        for row in tb.rows:
            cells = [c.text.strip() for c in row.cells]
            out.append(' | '.join(cells))
    return '\n'.join(out)[:limit]

targets = [
    os.path.join(root, '8、审计报告模版', '审计报告正文.docx'),
    os.path.join(root, '8、审计报告模版', '一般企业报表.docx'),
    os.path.join(root, '9、管理建议书模版', '管理建议书模板.docx'),
]
for p in targets:
    print('=' * 90)
    print(p)
    if HAS_DOCX:
        print(read_docx(p))
    else:
        print('NO python-docx')

# 复核关注点 xlsx
p = os.path.join(root, '11、复核要点、质量控制制度', '复核关注点.xlsx')
print('=' * 90)
print(p)
wb = openpyxl.load_workbook(p, read_only=True, data_only=True)
for ws in wb.worksheets:
    print('-- sheet:', ws.title)
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i >= 25:
            break
        vals = [str(v)[:20] if v is not None else '' for v in row[:10]]
        print('   ', ' | '.join(vals))
wb.close()

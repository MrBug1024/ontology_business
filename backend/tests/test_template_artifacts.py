from __future__ import annotations

import io
import json
from contextlib import nullcontext
import unittest
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from openpyxl import load_workbook
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool

from app.database import Base, get_db
from app.models import (
    Agent,
    BucketFile,
    BusinessScenario,
    Conversation,
    DataSource,
    Message,
    OntologyAction,
    OntologyBranch,
    OntologyEntity,
    OntologyRelease,
    OntologySnapshot,
    Tenant,
    User,
)
from app.routers import data_sources
from app.routers import scenarios as scenarios_router
from app.schemas import ActionIn
from app.services import (
    agent_capability_service,
    datasource_service,
    object_deletion_service,
    object_storage_service,
    permission_service,
    release_service,
    runtime_definition_service,
    template_artifact_service,
    workflow_service,
)
from app.services.auth_service import get_current_user, get_tenant_db
from tests.minio_fakes import FakeMinio


FIXTURES = Path(__file__).parent / "fixtures" / "artifact_templates"
TEST_MINIO_CONFIGURATION = object_storage_service.MinioConfiguration(
    endpoint="minio.example.test",
    access_key="access",
    secret_key="secret",
    bucket_name="ontology",
    prefix="ontology-business",
)
TEST_MINIO_SOURCE_CONFIG = {
    "storage_backend": "minio",
    "bucket_name": "ontology",
    "prefix": "ontology-business",
}


def _rewrite_zip_member(content: bytes, member_name: str, transform) -> bytes:
    output = io.BytesIO()
    with ZipFile(io.BytesIO(content), "r") as before, ZipFile(output, "w", ZIP_DEFLATED) as after:
        for info in before.infolist():
            payload = before.read(info)
            after.writestr(info, transform(payload) if info.filename == member_name else payload)
    return output.getvalue()


def _variables() -> dict:
    return {
        "project": {"name": "星河中心", "code": "PRJ-001"},
        "manager": {"name": "张三"},
        "report": {
            "date": "2026-08-22",
            "status": "正常",
            "summary": "结构施工按计划完成。",
            "risk": "连续降雨可能影响后续进度。",
        },
        "metrics": {
            "completion": 85,
            "contract_amount": 12_345_678.9,
            "budget": 1_000_000,
            "actual": 235_000,
        },
    }


def _docx_variables() -> dict:
    values = _variables()
    values["metrics"] = {
        "completion": values["metrics"]["completion"],
        "contract_amount": values["metrics"]["contract_amount"],
    }
    return values


class TemplateRenderingTests(unittest.TestCase):
    def test_template_paths_build_a_required_nested_schema_without_guessing_leaf_types(self) -> None:
        schema = template_artifact_service.merge_template_input_schema(
            {
                "type": "object",
                "properties": {
                    "metrics": {
                        "type": "object",
                        "properties": {"completion": {"type": "number"}},
                    }
                },
            },
            {"project.name", "metrics.completion", "lines.0.sku"},
        )
        self.assertEqual(schema["properties"]["metrics"]["properties"]["completion"]["type"], "number")
        self.assertNotIn("type", schema["properties"]["project"]["properties"]["name"])
        self.assertEqual(schema["properties"]["lines"]["type"], "array")
        self.assertEqual(schema["properties"]["lines"]["minItems"], 1)
        self.assertIn("sku", schema["properties"]["lines"]["items"]["required"])
        with self.assertRaisesRegex(template_artifact_service.TemplateArtifactError, "类型冲突"):
            template_artifact_service.merge_template_input_schema(
                {"type": "object", "properties": {"project": {"type": "string"}}},
                {"project.name"},
            )

    def test_markdown_keeps_markdown_and_resolves_nested_variables(self) -> None:
        source = FIXTURES / "项目周报模板.md"
        result = template_artifact_service.render_template(
            source.name,
            source.read_bytes(),
            _variables(),
            output_filename="星河中心周报.md",
        )
        text = result.content.decode("utf-8")
        self.assertEqual(result.format, "markdown")
        self.assertEqual(result.mime, template_artifact_service.MARKDOWN_MIME)
        self.assertEqual(result.filename, "星河中心周报.md")
        self.assertIn("# 星河中心项目周报", text)
        self.assertIn("完成率：85%", text)
        self.assertNotIn("{{", text)

        structured = template_artifact_service.render_template(
            "list.md",
            "首项={{items.0.name}}；完整={{items}}".encode(),
            {"items": [{"name": "塔楼"}, {"name": "裙房"}]},
        )
        self.assertIn("首项=塔楼", structured.content.decode())
        self.assertIn('[{"name":"塔楼"},{"name":"裙房"}]', structured.content.decode())

    def test_docx_replaces_split_runs_and_preserves_package_formatting(self) -> None:
        source = FIXTURES / "项目报告模板.docx"
        source_bytes = source.read_bytes()
        result = template_artifact_service.render_template(
            source.name,
            source_bytes,
            _variables(),
            output_filename="星河中心项目报告.docx",
        )
        self.assertEqual(result.format, "docx")
        self.assertEqual(result.mime, template_artifact_service.DOCX_MIME)
        self.assertTrue(result.content.startswith(b"PK"))

        with ZipFile(io.BytesIO(source_bytes)) as before, ZipFile(io.BytesIO(result.content)) as after:
            self.assertEqual(before.namelist(), after.namelist())
            self.assertEqual(before.read("word/styles.xml"), after.read("word/styles.xml"))
            self.assertEqual(before.read("word/settings.xml"), after.read("word/settings.xml"))
            rendered_xml = b"\n".join(
                after.read(name)
                for name in after.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            )
            self.assertNotIn(b"{{", rendered_xml)

        document = Document(io.BytesIO(result.content))
        body = "\n".join(paragraph.text for paragraph in document.paragraphs)
        table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        header_text = "\n".join(
            paragraph.text for section in document.sections for paragraph in section.header.paragraphs
        )
        self.assertIn("星河中心项目业务报告", body)
        self.assertIn("报告负责人：张三", body)
        self.assertIn("12,345,678.9".replace(",", ""), table_text.replace(",", ""))
        self.assertIn("项目编号：PRJ-001", header_text)

    def test_xlsx_replaces_values_but_preserves_styles_formulas_and_package(self) -> None:
        source = FIXTURES / "项目预算模板.xlsx"
        source_bytes = source.read_bytes()
        result = template_artifact_service.render_template(
            source.name,
            source_bytes,
            _variables(),
            output_filename="星河中心项目预算.xlsx",
        )
        self.assertEqual(result.format, "xlsx")
        self.assertEqual(result.mime, template_artifact_service.XLSX_MIME)
        self.assertTrue(result.content.startswith(b"PK"))

        with ZipFile(io.BytesIO(source_bytes)) as before, ZipFile(io.BytesIO(result.content)) as after:
            self.assertEqual(before.namelist(), after.namelist())
            self.assertEqual(before.read("xl/styles.xml"), after.read("xl/styles.xml"))

        before_book = load_workbook(io.BytesIO(source_bytes), data_only=False)
        after_book = load_workbook(io.BytesIO(result.content), data_only=False)
        before_sheet = before_book["项目预算"]
        after_sheet = after_book["项目预算"]
        self.assertEqual(after_sheet["A1"].value, "星河中心项目预算表")
        self.assertEqual(after_sheet["B3"].value, "PRJ-001")
        self.assertEqual(after_sheet["D3"].value, "张三")
        self.assertEqual(after_sheet["B4"].value, 1_000_000)
        self.assertEqual(after_sheet["B5"].value, 235_000)
        self.assertEqual(after_sheet["B6"].value, "=B4-B5")
        self.assertEqual(before_sheet["A1"].style_id, after_sheet["A1"].style_id)
        self.assertEqual(before_sheet["B6"].style_id, after_sheet["B6"].style_id)
        self.assertEqual(before_sheet.freeze_panes, after_sheet.freeze_panes)

    def test_xlsx_headers_and_footers_are_rendered_and_missing_values_block(self) -> None:
        source = (FIXTURES / "项目预算模板.xlsx").read_bytes()

        def add_header(raw: bytes, placeholder: bytes = b"{{project.name}}") -> bytes:
            prefix = b"x:" if b"</x:worksheet>" in raw else b""
            closing = b"</" + prefix + b"worksheet>"
            block = (
                b"<" + prefix + b"headerFooter><" + prefix + b"oddHeader>&amp;L"
                + placeholder
                + b"</" + prefix + b"oddHeader></" + prefix + b"headerFooter>"
            )
            return raw.replace(closing, block + closing)

        with_header = _rewrite_zip_member(
            source,
            "xl/worksheets/sheet1.xml",
            add_header,
        )
        result = template_artifact_service.render_template(
            "header.xlsx",
            with_header,
            _variables(),
        )
        with ZipFile(io.BytesIO(result.content)) as package:
            sheet_xml = package.read("xl/worksheets/sheet1.xml").decode("utf-8")
        self.assertIn("星河中心", sheet_xml)
        self.assertNotIn("{{project.name}}", sheet_xml)

        missing_header = _rewrite_zip_member(
            source,
            "xl/worksheets/sheet1.xml",
            lambda raw: add_header(raw, b"{{project.missing_header}}"),
        )
        with self.assertRaisesRegex(
            template_artifact_service.TemplateArtifactError,
            "缺少必填变量",
        ):
            template_artifact_service.render_template(
                "missing-header.xlsx",
                missing_header,
                _variables(),
            )

    def test_missing_values_wrong_extension_macro_and_traversal_fail_closed(self) -> None:
        markdown = (FIXTURES / "项目周报模板.md").read_bytes()
        with self.assertRaisesRegex(template_artifact_service.TemplateArtifactError, "缺少必填变量"):
            template_artifact_service.render_template("template.md", markdown, {})
        with self.assertRaisesRegex(template_artifact_service.TemplateArtifactError, "相同文件格式"):
            template_artifact_service.render_template(
                "template.md", markdown, _variables(), output_filename="report.docx"
            )
        with self.assertRaisesRegex(template_artifact_service.TemplateArtifactError, "宏"):
            template_artifact_service.render_template("unsafe.docm", b"not-office", {})
        with self.assertRaisesRegex(ValueError, "路径"):
            datasource_service.validate_bucket_filename("../report.md")

        docx = (FIXTURES / "项目报告模板.docx").read_bytes()
        with self.assertRaisesRegex(
            template_artifact_service.TemplateArtifactError, "真实内容类型|缺少工作表"
        ):
            template_artifact_service.render_template("renamed.xlsx", docx, _variables())

        source = io.BytesIO(docx)
        tampered = io.BytesIO()
        with ZipFile(source, "r") as before, ZipFile(tampered, "w", ZIP_DEFLATED) as after:
            for info in before.infolist():
                after.writestr(info, before.read(info))
            after.writestr("word/vbaProject.bin", b"macro")
        with self.assertRaisesRegex(template_artifact_service.TemplateArtifactError, "宏"):
            template_artifact_service.render_template(
                "unsafe.docx", tampered.getvalue(), _variables()
            )

    def test_zip_slip_executable_members_and_active_fields_fail_closed(self) -> None:
        docx = (FIXTURES / "项目报告模板.docx").read_bytes()
        for member_name in ("../escaped.xml", "word/payload.exe"):
            tampered = io.BytesIO()
            with ZipFile(io.BytesIO(docx), "r") as before, ZipFile(tampered, "w", ZIP_DEFLATED) as after:
                for info in before.infolist():
                    after.writestr(info, before.read(info))
                after.writestr(member_name, b"unsafe")
            with self.assertRaisesRegex(
                template_artifact_service.TemplateArtifactError,
                "不安全的压缩包路径|可执行文件成员",
            ):
                template_artifact_service.render_template(
                    "unsafe.docx",
                    tampered.getvalue(),
                    _variables(),
                )

        dde_docx = _rewrite_zip_member(
            docx,
            "word/document.xml",
            lambda raw: raw.replace(
                b"</w:body>",
                (
                    b'<w:p><w:fldSimple w:instr=" DDEAUTO cmd /c calc ">'
                    b"<w:r><w:t>unsafe</w:t></w:r></w:fldSimple></w:p></w:body>"
                ),
            ),
        )
        with self.assertRaisesRegex(
            template_artifact_service.TemplateArtifactError,
            "DDE|宏字段",
        ):
            template_artifact_service.render_template("dde.docx", dde_docx, _variables())

        xlsx = (FIXTURES / "项目预算模板.xlsx").read_bytes()
        dde_xlsx = _rewrite_zip_member(
            xlsx,
            "xl/worksheets/sheet1.xml",
            lambda raw: raw.replace(
                b"</x:sheetData>" if b"</x:sheetData>" in raw else b"</sheetData>",
                (
                    b"<x:row r=\"20\"><x:c r=\"A20\"><x:f>cmd|' /C calc'!A0</x:f>"
                    b"<x:v>0</x:v></x:c></x:row></x:sheetData>"
                    if b"</x:sheetData>" in raw
                    else b"<row r=\"20\"><c r=\"A20\"><f>cmd|' /C calc'!A0</f><v>0</v></c></row></sheetData>"
                ),
            ),
        )
        with self.assertRaisesRegex(
            template_artifact_service.TemplateArtifactError,
            "可执行或联网",
        ):
            template_artifact_service.render_template("dde.xlsx", dde_xlsx, _variables())


class TemplateActionAndDownloadTests(unittest.TestCase):
    @staticmethod
    def _fake_upload_claim(**kwargs):
        return SimpleNamespace(object_key=kwargs["object_key"])

    def setUp(self) -> None:
        self.engine = create_engine(
            "sqlite://",
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        self.Session = sessionmaker(bind=self.engine, expire_on_commit=False)
        Base.metadata.create_all(self.engine)
        self.minio = FakeMinio()
        self.configuration_patch = patch.object(
            object_storage_service,
            "configuration",
            return_value=TEST_MINIO_CONFIGURATION,
        )
        self.client_patch = patch.object(
            object_storage_service,
            "get_client",
            return_value=self.minio,
        )
        # Outbox transaction behavior has dedicated coverage. These template
        # tests use one StaticPool connection for HTTP and ORM assertions, so
        # a real upload-intent sub-session would incorrectly commit it.
        self.upload_intent_patch = patch.object(
            object_deletion_service,
            "_prepare_upload_intent",
            side_effect=self._fake_upload_claim,
        )
        self.heartbeat_patch = patch.object(
            object_deletion_service,
            "heartbeat_upload_intent",
            side_effect=lambda _claim: nullcontext(),
        )
        self.begin_upload_patch = patch.object(
            object_deletion_service,
            "begin_upload_put",
        )
        self.assert_upload_patch = patch.object(
            object_deletion_service,
            "assert_upload_active",
        )
        self.retain_upload_patch = patch.object(
            object_deletion_service,
            "retain_bucket_file_upload",
        )
        self.configuration_patch.start()
        self.client_patch.start()
        self.upload_intent_patch.start()
        self.heartbeat_patch.start()
        self.begin_upload_patch.start()
        self.assert_upload_patch.start()
        self.retain_upload_patch.start()
        self.tenant = Tenant(id="tenant-artifacts", name="附件测试租户")
        self.user = User(
            id="user-artifacts",
            tenant_id=self.tenant.id,
            email="artifact@example.test",
            password_hash="test-only",
            status="active",
        )
        self.foreign_tenant = Tenant(id="tenant-artifacts-foreign", name="其他附件租户")
        self.foreign_user = User(
            id="user-artifacts-foreign",
            tenant_id=self.foreign_tenant.id,
            email="artifact-foreign@example.test",
            password_hash="test-only",
            status="active",
        )
        self.scenario = BusinessScenario(
            id="scenario-artifacts", tenant_id=self.tenant.id, name="建筑项目"
        )
        self.entity = OntologyEntity(
            id="entity-artifacts", scenario_id=self.scenario.id, name="项目"
        )
        self.templates = DataSource(
            id="source-templates",
            tenant_id=self.tenant.id,
            scenario_id=self.scenario.id,
            name="业务模板",
            type="file_bucket",
            config=dict(TEST_MINIO_SOURCE_CONFIG),
        )
        self.outputs = DataSource(
            id="source-outputs",
            tenant_id=self.tenant.id,
            scenario_id=self.scenario.id,
            name="业务附件",
            type="file_bucket",
            config=dict(TEST_MINIO_SOURCE_CONFIG),
        )
        setup_db = self.Session()
        setup_db.add_all(
            [
                self.tenant,
                self.user,
                self.foreign_tenant,
                self.foreign_user,
                self.scenario,
                self.entity,
                self.templates,
                self.outputs,
            ]
        )
        setup_db.commit()
        permission_service.ensure_organization(
            setup_db, self.tenant.id, owner_user_id=self.user.id
        )
        permission_service.ensure_organization(
            setup_db,
            self.foreign_tenant.id,
            owner_user_id=self.foreign_user.id,
        )
        self.template_file = datasource_service.save_bucket_file(
            self.templates,
            "项目报告模板.docx",
            (FIXTURES / "项目报告模板.docx").read_bytes(),
            mime=template_artifact_service.DOCX_MIME,
        )
        self.template_file.status = "parsed"
        setup_db.add(self.template_file)
        setup_db.commit()
        setup_db.close()

        self.app = FastAPI()
        self.app.include_router(data_sources.router, prefix="/api")
        self.app.include_router(scenarios_router.router, prefix="/api")
        self.current_user_id = self.user.id
        self.current_tenant_id = self.tenant.id

        def override_user():
            return SimpleNamespace(id=self.current_user_id, tenant_id=self.current_tenant_id)

        def override_db():
            db = self.Session()
            db.info["tenant_id"] = self.current_tenant_id
            db.info["user_id"] = self.current_user_id
            try:
                yield db
            finally:
                db.close()

        self.app.dependency_overrides[get_current_user] = override_user
        self.app.dependency_overrides[get_db] = override_db
        self.app.dependency_overrides[get_tenant_db] = override_db
        self.client = TestClient(self.app)

    def tearDown(self) -> None:
        self.client.close()
        self.engine.dispose()
        self.retain_upload_patch.stop()
        self.assert_upload_patch.stop()
        self.begin_upload_patch.stop()
        self.heartbeat_patch.stop()
        self.upload_intent_patch.stop()
        self.client_patch.stop()
        self.configuration_patch.stop()

    def _db(self) -> Session:
        db = self.Session()
        db.info["tenant_id"] = self.tenant.id
        db.info["user_id"] = self.user.id
        return db

    def _create_template_action(self, db: Session, action_id: str) -> OntologyAction:
        payload = ActionIn(
            entity_id=self.entity.id,
            name="生成项目报告",
            input_schema={
                "type": "object",
                "properties": {
                    key: {"type": "object"}
                    for key in ("project", "manager", "report", "metrics")
                },
                "required": ["project", "manager", "report", "metrics"],
            },
            executor_type="template",
            executor_config={
                "template_file_id": self.template_file.id,
                "target_data_source_id": self.outputs.id,
                "output_filename": "{{project.name}}项目报告.docx",
                "template_sha256": "client-cannot-pin-this",
            },
            requires_confirmation=True,
            idempotency_required=True,
        )
        scenarios_router._validate_action_executor(db, self.scenario.id, payload)
        action = OntologyAction(
            id=action_id,
            scenario_id=self.scenario.id,
            **payload.model_dump(),
        )
        db.add(action)
        db.commit()
        return action

    def _release_current_definition(
        self,
        db: Session,
        *,
        branch_id: str,
        snapshot_id: str,
        release_id: str,
    ) -> None:
        branch = OntologyBranch(
            id=branch_id,
            tenant_id=self.tenant.id,
            scenario_id=self.scenario.id,
            name="main",
            created_by_user_id=self.user.id,
        )
        db.add(branch)
        db.flush()
        content = release_service.capture_snapshot_content(db, self.scenario)
        snapshot = OntologySnapshot(
            id=snapshot_id,
            tenant_id=self.tenant.id,
            scenario_id=self.scenario.id,
            branch_id=branch.id,
            kind="merge",
            content=content,
            content_hash=release_service.snapshot_hash(content),
            created_by_user_id=self.user.id,
        )
        db.add(snapshot)
        db.flush()
        branch.base_snapshot_id = snapshot.id
        branch.head_snapshot_id = snapshot.id
        db.add(
            OntologyRelease(
                id=release_id,
                tenant_id=self.tenant.id,
                scenario_id=self.scenario.id,
                branch_id=branch.id,
                snapshot_id=snapshot.id,
                environment="dev",
                status="released",
                created_by_user_id=self.user.id,
            )
        )
        db.commit()

    def test_template_action_pins_source_previews_then_generates_once(self) -> None:
        db = self._db()
        action = self._create_template_action(db, "action-template-report")
        pinned_hash = action.executor_config["template_sha256"]
        self.assertEqual(len(pinned_hash), 64)
        self.assertNotEqual(pinned_hash, "client-cannot-pin-this")
        self.assertIn("name", action.input_schema["properties"]["project"]["required"])
        self.assertNotIn(
            "type",
            action.input_schema["properties"]["project"]["properties"]["name"],
        )

        variables = _docx_variables()
        preview = workflow_service.execute_action(db, action, variables, dry_run=True)
        plan = preview["result"]["plan"]
        self.assertEqual(plan["artifact"]["filename"], "星河中心项目报告.docx")
        self.assertTrue(plan["side_effects_skipped"])
        self.assertEqual(
            db.query(BucketFile).filter_by(data_source_id=self.outputs.id).count(), 0
        )

        db.info["action_lineage_context"] = {
            "parent_action_log_id": preview["log_id"],
            "correlation_id": preview["correlation_id"],
        }
        executed = workflow_service.execute_action(
            db,
            action,
            variables,
            confirm=True,
            idempotency_key="project-report-001",
        )
        self.assertEqual(executed["status"], "success")
        artifact = executed["result"]["artifact"]
        self.assertEqual(artifact["format"], "docx")
        self.assertEqual(
            db.query(BucketFile).filter_by(data_source_id=self.outputs.id).count(), 1
        )

        replay = workflow_service.execute_action(
            db,
            action,
            variables,
            confirm=True,
            idempotency_key="a-different-client-key",
        )
        self.assertEqual(replay["status"], "idempotent_replay")
        self.assertEqual(
            db.query(BucketFile).filter_by(data_source_id=self.outputs.id).count(), 1
        )

        persisted = db.get(BucketFile, artifact["id"])
        self.assertEqual(persisted.generated_by_action_log_id, executed["log_id"])
        self.assertEqual(persisted.origin_template_file_id, self.template_file.id)
        self.assertEqual(persisted.content_sha256, artifact["sha256"])
        self.assertEqual(persisted.stored_path, persisted.object_url)
        self.assertIn(
            f"/files/{persisted.id}/uploads/",
            f"/{persisted.object_key}",
        )
        downloaded = self.client.get(artifact["download_url"])
        self.assertEqual(downloaded.status_code, 200, downloaded.text)
        self.assertEqual(downloaded.headers["content-type"], template_artifact_service.DOCX_MIME)
        self.assertIn("星河中心项目业务报告", "\n".join(
            paragraph.text for paragraph in Document(io.BytesIO(downloaded.content)).paragraphs
        ))

        original_size = persisted.size
        persisted.size += 1
        db.commit()
        corrupted = self.client.get(artifact["download_url"])
        self.assertEqual(corrupted.status_code, 409)
        self.assertIn("大小", corrupted.json()["detail"])
        persisted.size = original_size
        content = bytearray(
            datasource_service.read_bucket_file(persisted, self.outputs)[0]
        )
        content[-1] ^= 1
        self.minio.overwrite_object(
            persisted.bucket_name,
            persisted.object_key,
            bytes(content),
            version_id=persisted.object_version_id,
            preserve_etag=True,
        )
        db.commit()
        corrupted = self.client.get(artifact["download_url"])
        self.assertEqual(corrupted.status_code, 409)
        self.assertIn("哈希", corrupted.json()["detail"])
        db.info.pop("action_lineage_context", None)
        db.close()

    def test_template_commit_failure_can_retry_without_duplicate_file(self) -> None:
        db = self._db()
        action = self._create_template_action(db, "action-template-retry")
        with patch.object(db, "commit", side_effect=RuntimeError("simulated commit loss")):
            with self.assertRaisesRegex(RuntimeError, "commit loss"):
                workflow_service.execute_action(
                    db,
                    action,
                    _docx_variables(),
                    confirm=True,
                    idempotency_key="retry-after-rollback",
                )
        self.assertEqual(
            db.query(BucketFile).filter_by(data_source_id=self.outputs.id).count(), 0
        )
        result = workflow_service.execute_action(
            db,
            action,
            _docx_variables(),
            confirm=True,
            idempotency_key="retry-after-rollback",
        )
        self.assertEqual(result["status"], "success")
        persisted = db.get(BucketFile, result["result"]["artifact"]["id"])
        self.assertTrue(
            datasource_service.read_bucket_file(persisted, self.outputs)[0].startswith(
                b"PK"
            )
        )
        scoped_objects = object_storage_service.list_objects(
            persisted.bucket_name,
            f"ontology-business/tenants/{self.tenant.id}/scenarios/{self.scenario.id}/"
            f"data-sources/{self.outputs.id}/files/{persisted.id}",
        )
        # The failed PUT uses a different generation and is later reclaimed by
        # the outbox. The durable metadata must point only to the retry output.
        self.assertIn(
            persisted.object_key,
            [item.object_key for item in scoped_objects],
        )
        db.close()

    def test_agent_confirmation_revalidates_scope_conversation_and_download_acl(self) -> None:
        db = self._db()
        action = self._create_template_action(db, "action-template-agent")
        scope = agent_capability_service.explicit_empty_scope()
        scope["actions"] = {"mode": "explicit", "selected_ids": [action.id]}
        agent = Agent(
            id="agent-template-confirm",
            tenant_id=self.tenant.id,
            scenario_id=self.scenario.id,
            name="模板附件 Agent",
            capability_scope=scope,
        )
        conversation = Conversation(
            id="conversation-template-confirm",
            agent=agent,
            created_by_user_id=self.user.id,
            title="生成项目报告",
        )
        message = Message(
            id="message-template-confirm",
            conversation=conversation,
            role="assistant",
            content="报告预演已完成。",
            stream_finalized=True,
            tool_calls=[{"id": "tool-template", "name": "execute_action", "arguments": {}}],
            tool_results=[],
        )
        other_conversation = Conversation(
            id="conversation-template-other",
            agent=agent,
            created_by_user_id=self.user.id,
            title="其他对话",
        )
        other_message = Message(
            id="message-template-other",
            conversation=other_conversation,
            role="assistant",
            content="不应被确认结果修改。",
            stream_finalized=True,
            tool_calls=[],
            tool_results=[],
        )
        db.add_all([agent, conversation, message, other_conversation, other_message])
        # This confirmation case validates a released Action. The generic
        # fixture entity has no business key, so mark it abstract before
        # capturing the otherwise valid release snapshot.
        released_entity = db.get(OntologyEntity, self.entity.id)
        assert released_entity is not None
        released_entity.is_abstract = True
        db.commit()
        self._release_current_definition(
            db,
            branch_id="branch-template-agent",
            snapshot_id="snapshot-template-agent",
            release_id="release-template-agent",
        )
        definition = runtime_definition_service.resolve_execution(
            db,
            self.scenario,
            environment="dev",
        )
        released_action = runtime_definition_service.resolve_resource(
            definition,
            "action",
            action.id,
        )
        db.info["action_audit_context"] = {"agent_id": agent.id}
        db.info["llm_trace_context"] = {
            "assistant_message_id": message.id,
            "correlation_id": "template-confirm-correlation",
        }
        preview = workflow_service.execute_action(
            db,
            released_action,
            _docx_variables(),
            dry_run=True,
            runtime_environment="dev",
            runtime_definition=definition,
        )
        db.info.pop("action_audit_context", None)
        db.info.pop("llm_trace_context", None)
        message.tool_results = [{
            "id": "tool-template",
            "name": "execute_action",
            "result": json.dumps(preview, ensure_ascii=False),
        }]
        db.commit()

        payload = {
            "params": _docx_variables(),
            "dry_run": False,
            "confirm": True,
            "idempotency_key": "agent-template-confirm",
            "preview_log_id": preview["log_id"],
            "correlation_id": preview["correlation_id"],
            "expected_environment": preview["environment"],
            "expected_definition_snapshot_id": preview["definition_snapshot_id"],
            "expected_release_id": preview["release_id"],
            "expected_definition_hash": preview["definition_hash"],
        }

        saved_tool_results = list(message.tool_results)
        message.tool_results = []
        db.commit()
        missing_message_preview = self.client.post(
            f"/api/scenarios/actions/{action.id}/execute",
            json=payload,
        )
        self.assertEqual(missing_message_preview.status_code, 409, missing_message_preview.text)
        self.assertIn("对话中未找到", missing_message_preview.json()["detail"])
        message.tool_results = saved_tool_results
        db.commit()

        agent.capability_scope = agent_capability_service.explicit_empty_scope()
        db.commit()
        revoked = self.client.post(
            f"/api/scenarios/actions/{action.id}/execute",
            json=payload,
        )
        self.assertEqual(revoked.status_code, 409, revoked.text)
        self.assertIn("Agent", revoked.json()["detail"])

        agent.capability_scope = scope
        db.commit()
        self.current_user_id = self.foreign_user.id
        self.current_tenant_id = self.foreign_tenant.id
        foreign_confirmation = self.client.post(
            f"/api/scenarios/actions/{action.id}/execute",
            json=payload,
        )
        self.assertIn(foreign_confirmation.status_code, {403, 404})

        self.current_user_id = self.user.id
        self.current_tenant_id = self.tenant.id
        confirmed = self.client.post(
            f"/api/scenarios/actions/{action.id}/execute",
            json=payload,
        )
        self.assertEqual(confirmed.status_code, 200, confirmed.text)
        result = confirmed.json()
        self.assertEqual(result["status"], "success")
        artifact = result["result"]["artifact"]
        db.expire_all()
        stored_message = db.get(Message, message.id)
        stored_result = json.loads(stored_message.tool_results[0]["result"])
        self.assertEqual(stored_result["result"]["artifact"]["id"], artifact["id"])
        self.assertEqual(db.get(Message, other_message.id).tool_results, [])

        self.current_user_id = self.foreign_user.id
        self.current_tenant_id = self.foreign_tenant.id
        forbidden_download = self.client.get(artifact["download_url"])
        self.assertIn(forbidden_download.status_code, {403, 404})
        self.current_user_id = self.user.id
        self.current_tenant_id = self.tenant.id
        allowed_download = self.client.get(artifact["download_url"])
        self.assertEqual(allowed_download.status_code, 200, allowed_download.text)
        self.assertEqual(
            allowed_download.headers["content-type"],
            template_artifact_service.DOCX_MIME,
        )
        db.close()

    def test_stable_retry_uses_a_fresh_minio_generation(self) -> None:
        db = self._db()
        action = self._create_template_action(db, "action-template-crash")
        scoped_key = workflow_service._scoped_idempotency_key("crash-retry", "dev")
        stable_id = workflow_service._stable_template_execution_id(
            action,
            parent_action_log_id=None,
            scoped_idempotency_key=scoped_key,
            environment="dev",
        )
        orphan = datasource_service.save_bucket_file(
            self.outputs,
            "星河中心项目报告.docx",
            b"interrupted-write-placeholder",
            stable_file_id=stable_id,
        )
        result = workflow_service.execute_action(
            db,
            action,
            _docx_variables(),
            confirm=True,
            idempotency_key="crash-retry",
        )
        self.assertEqual(result["result"]["artifact"]["id"], stable_id)
        persisted = db.get(BucketFile, stable_id)
        self.assertNotEqual(orphan.object_key, persisted.object_key)
        self.assertEqual(
            datasource_service.read_bucket_file(orphan, self.outputs)[0],
            b"interrupted-write-placeholder",
        )
        self.assertTrue(
            datasource_service.read_bucket_file(persisted, self.outputs)[0].startswith(
                b"PK"
            )
        )
        db.close()

    def test_bucket_paths_and_duplicate_names_are_safe(self) -> None:
        first = datasource_service.save_bucket_file(self.outputs, "报告.md", b"one")
        second = datasource_service.save_bucket_file(self.outputs, "报告.md", b"two")
        self.assertEqual(first.filename, "报告.md")
        self.assertEqual(second.filename, "报告.md")
        self.assertNotEqual(first.object_key, second.object_key)
        self.assertEqual(
            datasource_service.read_bucket_file(second, self.outputs)[0],
            b"two",
        )
        first.mime = "application/octet-stream"
        _content, _size, canonical_mime = datasource_service.read_bucket_file(
            first,
            self.outputs,
        )
        self.assertEqual(canonical_mime, template_artifact_service.MARKDOWN_MIME)
        first.mime = template_artifact_service.DOCX_MIME
        with self.assertRaisesRegex(ValueError, "MIME"):
            datasource_service.read_bucket_file(first, self.outputs)

        escaped_key = (
            "ontology-business/tenants/other-tenant/scenarios/"
            f"{self.scenario.id}/data-sources/{self.outputs.id}/files/escaped/"
            f"uploads/{'a' * 32}/outside.md"
        )
        escaped_url = object_storage_service.stable_object_url("ontology", escaped_key)
        escaped = BucketFile(
            id="escaped",
            data_source_id=self.outputs.id,
            filename="outside.md",
            stored_path=escaped_url,
            storage_provider="minio",
            bucket_name="ontology",
            object_key=escaped_key,
            object_url=escaped_url,
            size=6,
            mime=template_artifact_service.MARKDOWN_MIME,
        )
        with self.assertRaisesRegex(ValueError, "不属于指定文件桶作用域"):
            datasource_service.read_bucket_file(escaped, self.outputs)


if __name__ == "__main__":
    unittest.main()
    Message,

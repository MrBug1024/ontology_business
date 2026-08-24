"""Rebuild the checked-in DOCX regression template with deterministic content."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("项目报告模板.docx")

document = Document()
section = document.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

normal = document.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(10.5)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
header_run = header.add_run("项目编号：{{project.code}}")
header_run.font.name = "Arial"
header_run.font.size = Pt(9)
header_run.font.color.rgb = RGBColor(91, 101, 115)

title = document.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("{{project.name}}项目业务报告")
title_run.bold = True
title_run.font.name = "Arial"
title_run.font.size = Pt(20)
title_run.font.color.rgb = RGBColor(30, 64, 175)

lead = document.add_paragraph()
lead.add_run("报告负责人：").bold = True
# Deliberately split one placeholder across runs to guard Word run boundaries.
lead.add_run("{{manager.")
lead.add_run("name}}")

table = document.add_table(rows=4, cols=2)
table.style = "Light Shading Accent 1"
table.autofit = False
rows = [
    ("报告日期", "{{report.date}}"),
    ("当前状态", "{{report.status}}"),
    ("合同金额", "{{metrics.contract_amount}}"),
    ("完成率", "{{metrics.completion}}%"),
]
for row, values in zip(table.rows, rows):
    row.cells[0].width = Inches(1.6)
    row.cells[1].width = Inches(5.2)
    row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row.cells[0].text = values[0]
    row.cells[1].text = values[1]

document.add_heading("业务摘要", level=1)
document.add_paragraph("{{report.summary}}")
document.add_heading("风险提示", level=1)
document.add_paragraph("{{report.risk}}")

document.save(OUTPUT)


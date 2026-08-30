from __future__ import annotations

from contextlib import ExitStack, nullcontext
from io import BytesIO
import hashlib
import json
from types import SimpleNamespace
import unittest
from unittest.mock import patch

from fastapi import FastAPI
from fastapi.testclient import TestClient
from sqlalchemy import create_engine, func, select
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.database import Base, get_db
from app.models import (
    BucketFile,
    DataAsset,
    DataAssetVersion,
    DataSource,
    DatasetVersion,
    LogicalDataset,
    ScenarioDatasetBinding,
    SemanticMapping,
    Tenant,
    User,
)
from app.routers import catalog
from app.services import (
    catalog_ingestion_service,
    datasource_service,
    object_deletion_service,
    object_storage_service,
    permission_service,
)
from app.services.auth_service import get_current_user, get_tenant_db


class CatalogManagedUploadTests(unittest.TestCase):
    def setUp(self) -> None:
        self.engine = create_engine(
            "sqlite://",
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        self.Session = sessionmaker(
            bind=self.engine,
            autoflush=False,
            expire_on_commit=False,
        )
        Base.metadata.create_all(self.engine)
        with self.Session() as db:
            tenant = Tenant(id="tenant-upload", name="Upload tenant")
            user = User(
                id="user-upload",
                tenant_id=tenant.id,
                email="upload@example.test",
                password_hash="test-only",
                status="active",
            )
            other_tenant = Tenant(id="tenant-upload-other", name="Other tenant")
            other_user = User(
                id="user-upload-other",
                tenant_id=other_tenant.id,
                email="upload-other@example.test",
                password_hash="test-only",
                status="active",
            )
            source = DataSource(
                id="managed-bucket",
                tenant_id=tenant.id,
                name="Managed physical storage",
                type="file_bucket",
                config={
                    "storage_backend": "minio",
                    "bucket_name": "catalog-test",
                    "prefix": "platform",
                },
            )
            other_source = DataSource(
                id="other-managed-bucket",
                tenant_id=other_tenant.id,
                name="Foreign managed storage",
                type="file_bucket",
                config={
                    "storage_backend": "minio",
                    "bucket_name": "catalog-test",
                    "prefix": "platform",
                },
            )
            db.add_all([tenant, user, other_tenant, other_user, source, other_source])
            db.commit()
            permission_service.ensure_organization(
                db, tenant.id, owner_user_id=user.id
            )
            permission_service.ensure_organization(
                db, other_tenant.id, owner_user_id=other_user.id
            )
            db.commit()

        self.app = FastAPI()
        self.app.include_router(catalog.router, prefix="/api")

        def override_user():
            return SimpleNamespace(id="user-upload", tenant_id="tenant-upload")

        def override_db():
            db = self.Session()
            db.info["tenant_id"] = "tenant-upload"
            db.info["user_id"] = "user-upload"
            try:
                yield db
            finally:
                db.close()

        self.app.dependency_overrides[get_current_user] = override_user
        self.app.dependency_overrides[get_db] = override_db
        self.app.dependency_overrides[get_tenant_db] = override_db
        self.client = TestClient(self.app)

        self.saved_calls: list[dict[str, object]] = []
        self.patches = ExitStack()
        configuration = SimpleNamespace(
            bucket_name="catalog-test",
            prefix="platform",
        )
        self.patches.enter_context(
            patch.object(
                object_storage_service,
                "require_configuration",
                return_value=configuration,
            )
        )
        self.patches.enter_context(
            patch.object(object_storage_service, "ensure_bucket")
        )

        def prepare_claim(source, file_id, filename):
            object_key = datasource_service.build_bucket_object_key(
                source,
                file_id,
                filename,
                upload_id="a" * 32,
            )
            return SimpleNamespace(
                object_key=object_key,
                bucket_name="catalog-test",
                provider="minio",
            )

        def save_file(
            source,
            filename,
            content,
            *,
            mime=None,
            stable_file_id=None,
            upload_object_key=None,
        ):
            self.saved_calls.append(
                {
                    "source_id": source.id,
                    "filename": filename,
                    "content": content,
                    "object_key": upload_object_key,
                }
            )
            digest = hashlib.sha256(content).hexdigest()
            return BucketFile(
                id=stable_file_id,
                data_source_id=source.id,
                filename=filename,
                stored_path=f"minio://catalog-test/{upload_object_key}",
                storage_provider="minio",
                bucket_name="catalog-test",
                object_key=upload_object_key,
                object_version_id="v1",
                etag=digest,
                object_url=f"minio://catalog-test/{upload_object_key}",
                size=len(content),
                mime=mime or "application/octet-stream",
                content_sha256=digest,
            )

        self.patches.enter_context(
            patch.object(
                object_deletion_service,
                "prepare_bucket_file_upload",
                side_effect=prepare_claim,
            )
        )
        self.patches.enter_context(
            patch.object(
                object_deletion_service,
                "heartbeat_upload_intent",
                side_effect=lambda _claim: nullcontext(SimpleNamespace()),
            )
        )
        self.patches.enter_context(
            patch.object(object_deletion_service, "begin_upload_put")
        )
        self.patches.enter_context(
            patch.object(object_deletion_service, "assert_upload_active")
        )
        self.patches.enter_context(
            patch.object(object_deletion_service, "retain_bucket_file_upload")
        )
        self.patches.enter_context(
            patch.object(
                datasource_service,
                "save_bucket_file",
                side_effect=save_file,
            )
        )

    def tearDown(self) -> None:
        self.patches.close()
        self.client.close()
        self.engine.dispose()

    def _upload(
        self,
        content: bytes,
        filename: str,
        media_type: str,
        **fields,
    ):
        data = {
            "file_bucket_id": "managed-bucket",
            "purpose": "managed_asset",
            **fields,
        }
        return self.client.post(
            "/api/catalog/uploads",
            data=data,
            files={"file": (filename, content, media_type)},
        )

    def test_csv_profile_is_schema_only_and_content_hash_retry_is_idempotent(self) -> None:
        content = (
            "record_id,amount,active,occurred_on,note\n"
            "A-1,12.50,true,2026-08-01,private-alpha\n"
            "A-2,,false,2026-08-02,private-beta\n"
        ).encode()
        first = self._upload(
            content,
            "records.csv",
            "text/csv",
            asset_key="generic.records.upload",
            labels=json.dumps({"source_class": "user_upload"}),
        )
        self.assertEqual(first.status_code, 201, first.text)
        body = first.json()
        self.assertTrue(body["created"])
        self.assertFalse(body["temporary"])
        table = body["version"]["profile"]["tables"][0]
        self.assertEqual(table["sample_row_count"], 2)
        self.assertEqual(
            [column["name"] for column in table["columns"]],
            ["record_id", "amount", "active", "occurred_on", "note"],
        )
        self.assertEqual(table["columns"][1]["logical_type"], "number")
        self.assertTrue(table["columns"][1]["nullable"])
        serialized = json.dumps(body, ensure_ascii=False)
        self.assertNotIn("private-alpha", serialized)
        self.assertNotIn("private-beta", serialized)
        for physical_name in ("bucket_name", "object_key", "data_source_id"):
            self.assertNotIn(physical_name, serialized)

        listed = self.client.get(
            f"/api/catalog/assets/{body['asset']['id']}/versions"
        )
        self.assertEqual(listed.status_code, 200, listed.text)
        listed_serialized = json.dumps(listed.json())
        self.assertNotIn("bucket_file_id", listed_serialized)
        self.assertNotIn("source_locator", listed_serialized)

        second = self._upload(
            content,
            "records.csv",
            "text/csv",
            asset_key="generic.records.upload",
            labels=json.dumps({"source_class": "user_upload"}),
        )
        self.assertEqual(second.status_code, 201, second.text)
        self.assertFalse(second.json()["created"])
        self.assertEqual(second.json()["asset"]["id"], body["asset"]["id"])
        self.assertEqual(second.json()["version"]["id"], body["version"]["id"])
        self.assertEqual(len(self.saved_calls), 1)
        with self.Session() as db:
            self.assertEqual(db.scalar(select(func.count(DataAsset.id))), 1)
            self.assertEqual(db.scalar(select(func.count(DataAssetVersion.id))), 1)
            self.assertEqual(db.scalar(select(func.count(BucketFile.id))), 1)
            self.assertEqual(db.scalar(select(func.count(LogicalDataset.id))), 0)
            self.assertEqual(db.scalar(select(func.count(DatasetVersion.id))), 0)
            self.assertEqual(db.scalar(select(func.count(ScenarioDatasetBinding.id))), 0)
            self.assertEqual(db.scalar(select(func.count(SemanticMapping.id))), 0)

    def test_xlsx_profile_covers_sheets_headers_and_logical_types(self) -> None:
        from openpyxl import Workbook

        output = BytesIO()
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Items"
        sheet.append(["item", "count", "enabled"])
        sheet.append(["one", 2, True])
        sheet.append(["two", None, False])
        workbook.save(output)
        response = self._upload(
            output.getvalue(),
            "items.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        self.assertEqual(response.status_code, 201, response.text)
        table = response.json()["version"]["profile"]["tables"][0]
        self.assertEqual(table["name"], "Items")
        self.assertEqual(table["columns"][1]["logical_type"], "integer")
        self.assertTrue(table["columns"][1]["nullable"])
        self.assertEqual(table["columns"][2]["logical_type"], "boolean")

    def test_document_profile_records_scale_without_persisting_parsed_text(self) -> None:
        from docx import Document

        output = BytesIO()
        document = Document()
        document.add_heading("Generic document", level=1)
        document.add_paragraph("confidential paragraph")
        document.save(output)
        response = self._upload(
            output.getvalue(),
            "requirements.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertEqual(response.status_code, 201, response.text)
        profile = response.json()["version"]["profile"]
        self.assertEqual(profile["category"], "document")
        self.assertEqual(profile["parser"]["status"], "parsed")
        self.assertGreater(profile["text"]["character_count"], 0)
        self.assertNotIn("confidential paragraph", json.dumps(profile))
        with self.Session() as db:
            bucket_file = db.scalar(select(BucketFile))
            self.assertEqual(bucket_file.parsed_text, "")

    def test_temporary_attachment_has_expiry_and_cannot_be_promoted_in_place(self) -> None:
        content = b"key,value\nA,1\n"
        temporary = self._upload(
            content,
            "input.csv",
            "text/csv",
            purpose="invocation_attachment",
            asset_key="request.input.attachment",
            expires_in_seconds="600",
        )
        self.assertEqual(temporary.status_code, 201, temporary.text)
        body = temporary.json()
        self.assertTrue(body["temporary"])
        self.assertIsNotNone(body["expires_at"])
        self.assertEqual(
            body["version"]["lifecycle"]["promotion_policy"],
            "explicit_copy_required",
        )
        self.assertFalse(body["version"]["lifecycle"]["auto_promote"])

        promoted = self._upload(
            content,
            "input.csv",
            "text/csv",
            purpose="managed_asset",
            asset_key="request.input.attachment",
        )
        self.assertEqual(promoted.status_code, 400, promoted.text)
        replaced = self._upload(
            b"key,value\nB,2\n",
            "input.csv",
            "text/csv",
            purpose="invocation_attachment",
            asset_key="request.input.attachment",
            expires_in_seconds="600",
        )
        self.assertEqual(replaced.status_code, 400, replaced.text)
        with self.Session() as db:
            asset = db.scalar(
                select(DataAsset).where(DataAsset.key == "request.input.attachment")
            )
            self.assertEqual(asset.labels["catalog_purpose"], "invocation_attachment")
            self.assertTrue(asset.labels["temporary"])
            self.assertEqual(db.scalar(select(func.count(LogicalDataset.id))), 0)
            self.assertEqual(db.scalar(select(func.count(ScenarioDatasetBinding.id))), 0)

    def test_cross_tenant_physical_fields_credentials_and_bad_files_are_rejected(self) -> None:
        cross_tenant = self._upload(
            b"a,b\n1,2\n",
            "records.csv",
            "text/csv",
            file_bucket_id="other-managed-bucket",
        )
        self.assertEqual(cross_tenant.status_code, 400, cross_tenant.text)

        physical = self._upload(
            b"a,b\n1,2\n",
            "records.csv",
            "text/csv",
            bucket_name="attacker-bucket",
        )
        self.assertEqual(physical.status_code, 400, physical.text)

        credentials = self._upload(
            b"a,b\n1,2\n",
            "records.csv",
            "text/csv",
            labels=json.dumps({"api_key": "do-not-store"}),
        )
        self.assertEqual(credentials.status_code, 400, credentials.text)

        hidden_locator = self._upload(
            b"a,b\n1,2\n",
            "records.csv",
            "text/csv",
            labels=json.dumps({"storage": {"object_key": "chosen/by/client"}}),
        )
        self.assertEqual(hidden_locator.status_code, 400, hidden_locator.text)

        traversal = self._upload(
            b"a,b\n1,2\n",
            "../records.csv",
            "text/csv",
        )
        self.assertEqual(traversal.status_code, 400, traversal.text)

        unsupported = self._upload(b"MZ", "program.exe", "application/octet-stream")
        self.assertEqual(unsupported.status_code, 400, unsupported.text)

        mime_mismatch = self._upload(
            b"a,b\n1,2\n",
            "records.csv",
            "image/png",
        )
        self.assertEqual(mime_mismatch.status_code, 400, mime_mismatch.text)
        self.assertEqual(len(self.saved_calls), 0)

    def test_upload_size_limit_is_enforced_before_storage(self) -> None:
        with patch.object(
            catalog,
            "get_settings",
            return_value=SimpleNamespace(max_upload_bytes=4),
        ):
            response = self._upload(b"12345", "small.txt", "text/plain")
        self.assertEqual(response.status_code, 413, response.text)
        self.assertEqual(len(self.saved_calls), 0)

    def test_lost_upload_lease_rolls_back_metadata_and_schedules_cleanup(self) -> None:
        with (
            patch.object(
                object_deletion_service,
                "assert_upload_active",
                side_effect=object_deletion_service.UploadIntentLeaseLostError(
                    "test lease lost"
                ),
            ),
            patch.object(
                object_deletion_service,
                "schedule_abandoned_upload_best_effort",
            ) as cleanup,
        ):
            response = self._upload(
                b"a,b\n1,2\n",
                "records.csv",
                "text/csv",
            )
        self.assertEqual(response.status_code, 503, response.text)
        cleanup.assert_called_once()
        with self.Session() as db:
            self.assertEqual(db.scalar(select(func.count(DataAsset.id))), 0)
            self.assertEqual(db.scalar(select(func.count(DataAssetVersion.id))), 0)
            self.assertEqual(db.scalar(select(func.count(BucketFile.id))), 0)

    def test_post_upload_integrity_conflict_schedules_object_cleanup(self) -> None:
        with (
            patch.object(
                catalog_ingestion_service,
                "register_prepared_version",
                side_effect=IntegrityError("insert", {}, RuntimeError("race")),
            ),
            patch.object(
                object_deletion_service,
                "schedule_abandoned_upload_best_effort",
            ) as cleanup,
        ):
            response = self._upload(
                b"a,b\n1,2\n",
                "records.csv",
                "text/csv",
            )

        self.assertEqual(response.status_code, 409, response.text)
        cleanup.assert_called_once()


if __name__ == "__main__":
    unittest.main()
